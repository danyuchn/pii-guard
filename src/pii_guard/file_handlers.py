"""Bounded, isolated multi-format file reading and writing.

The format libraries used by this module process untrusted user documents.
They therefore run in a fresh ``spawn`` child process. The parent process only
performs a bounded, race-checked byte read and accepts a validated JSON
response. Workers never send exception text, paths, or library output back to
the caller.

PDF output is deliberately rasterised after redaction. This removes the
source document's hidden text layer, metadata, and attachments. If a mapped
value cannot be located or the output cannot be validated, no PDF is written.
"""

from __future__ import annotations

import base64
import io
import json
import logging
import multiprocessing
import os
import stat
import sys
import threading
import time
import unicodedata
from collections.abc import Mapping, Sequence
from contextlib import contextmanager, redirect_stderr, redirect_stdout
from dataclasses import dataclass, field
from multiprocessing import BufferTooShort
from multiprocessing.process import BaseProcess
from pathlib import Path
from typing import Final, NoReturn, cast

from pii_guard._compat import is_reparse_point

logger = logging.getLogger(__name__)

PLAIN_TEXT_EXTENSIONS: set[str] = {
    ".txt",
    ".csv",
    ".tsv",
    ".log",
    ".md",
    ".dat",
    ".json",
    ".xml",
    ".html",
    ".htm",
    ".yaml",
    ".yml",
}
EXCEL_EXTENSIONS: set[str] = {".xlsx"}
WORD_EXTENSIONS: set[str] = {".docx"}
PDF_EXTENSIONS: set[str] = {".pdf"}
SUPPORTED_EXTENSIONS: set[str] = (
    PLAIN_TEXT_EXTENSIONS | EXCEL_EXTENSIONS | WORD_EXTENSIONS | PDF_EXTENSIONS
)

# This limit bounds both the bytes handed to a third-party parser and the
# maximum amount of data that can be returned over the worker pipe.
MAX_FILE_BYTES: Final[int] = 16 * 1024 * 1024
MAX_FILE_RESPONSE_BYTES: Final[int] = 64 * 1024 * 1024
FILE_PARSE_TIMEOUT_SECONDS: Final[float] = 15.0
FILE_PARSE_GRACE_SECONDS: Final[float] = 0.5
FILE_WORKER_CPU_SECONDS: Final[int] = 10
FILE_WORKER_ADDRESS_SPACE_BYTES: Final[int] = 512 * 1024 * 1024
PDF_MAX_PAGES: Final[int] = 50
PDF_RENDER_SCALE: Final[float] = 2.0
# pypdfium2 renders a page's CropBox. If it ever clips that CropBox back to
# the page's MediaBox (e.g. a malformed box), the rendered pixel size will
# disagree with pdfplumber's CropBox size by more than float rounding noise.
PDF_CROPBOX_RENDER_TOLERANCE_POINTS: Final[float] = 1.0

_FILE_ERROR_MESSAGES: Final[dict[str, str]] = {
    "FILE_NOT_FOUND": "Input file is not available.",
    "FILE_NOT_REGULAR": "Input file is not a regular file.",
    "FILE_UNSUPPORTED": "Input file format is not supported.",
    "FILE_TOO_LARGE": "Input file exceeds the safety size limit.",
    "FILE_CHANGED": "Input file changed during processing.",
    "FILE_OPEN_FAILED": "Input file could not be opened safely.",
    "FILE_NOT_UTF8": "Input file must be UTF-8 text.",
    "FORMAT_UNAVAILABLE": "Required file format support is unavailable.",
    "FILE_MALFORMED": "Input file could not be parsed safely.",
    "FILE_PARSE_TIMEOUT": "File parsing timed out safely.",
    "FILE_PARSE_CRASHED": "File parser stopped unexpectedly.",
    "FILE_PARSE_RESOURCE_LIMIT": "File parser exceeded the safety resource limit.",
    "FILE_RESPONSE_TOO_LARGE": "File parser returned an unsafe response.",
    "FILE_RESPONSE_INVALID": "File parser returned an invalid response.",
    "FILE_IPC_FAILED": "File parser communication failed safely.",
    "FILE_WRITE_FAILED": "The requested output could not be written safely.",
    "PDF_REDACTION_UNRESOLVED": "PDF redaction could not be completed safely.",
    "PDF_RENDER_UNAVAILABLE": "PDF rendering support is unavailable.",
    "PDF_OUTPUT_INVALID": "Generated PDF failed safety validation.",
}


class FileHandlerError(Exception):
    """A fixed, caller-safe file handling error."""

    def __init__(self, code: str, message: str | None = None) -> None:
        del message
        safe_code = code if code in _FILE_ERROR_MESSAGES else "FILE_MALFORMED"
        super().__init__(_FILE_ERROR_MESSAGES[safe_code])
        self.code = safe_code
        self.message = _FILE_ERROR_MESSAGES[safe_code]

    def __str__(self) -> str:
        return f"{self.code}: {self.message}"


@dataclass
class FileContent:
    """Extracted file text and private bytes needed for write-back."""

    text: str
    file_type: str  # "plain", "excel", "docx", "pdf"
    source_path: str = ""
    cells: list[tuple[str, str]] = field(default_factory=list)
    # Kept private from repr and never returned by parser IPC.
    source_bytes: bytes = field(default=b"", repr=False)


def _raise_file_error(code: str) -> NoReturn:
    """Raise a fixed public-safe error without exception context."""

    safe_code = code if code in _FILE_ERROR_MESSAGES else "FILE_MALFORMED"
    raise FileHandlerError(safe_code) from None


def _check_import(package: str, pip_name: str) -> None:
    """Raise a fixed error if an optional format package is missing."""

    del pip_name  # Kept for compatibility with the old private helper API.
    try:
        __import__(package)
    except (ImportError, ModuleNotFoundError):
        _raise_file_error("FORMAT_UNAVAILABLE")


def _read_source_bytes(path: Path) -> bytes:
    """Read a bounded regular file with a symlink and TOCTOU check."""

    try:
        info = path.lstat()
    except FileNotFoundError:
        _raise_file_error("FILE_NOT_FOUND")
    except (OSError, ValueError):
        _raise_file_error("FILE_OPEN_FAILED")

    current_path = Path(path.anchor or ".")
    parts = path.parts[1:] if path.anchor else path.parts
    for component in parts:
        current_path /= component
        try:
            component_info = current_path.lstat()
        except FileNotFoundError:
            break
        except (OSError, ValueError):
            _raise_file_error("FILE_OPEN_FAILED")
        if is_reparse_point(component_info):
            _raise_file_error("FILE_NOT_REGULAR")

    if is_reparse_point(info) or not stat.S_ISREG(info.st_mode):
        _raise_file_error("FILE_NOT_REGULAR")
    if info.st_size > MAX_FILE_BYTES:
        _raise_file_error("FILE_TOO_LARGE")

    flags = os.O_RDONLY | getattr(os, "O_NOFOLLOW", 0)
    try:
        descriptor = os.open(path, flags)
    except (OSError, ValueError):
        _raise_file_error("FILE_OPEN_FAILED")

    chunks: list[bytes] = []
    total = 0
    try:
        current = os.fstat(descriptor)
        if not stat.S_ISREG(current.st_mode):
            _raise_file_error("FILE_NOT_REGULAR")
        if current.st_ino != info.st_ino or current.st_dev != info.st_dev:
            _raise_file_error("FILE_CHANGED")
        while total <= MAX_FILE_BYTES:
            chunk = os.read(descriptor, min(1024 * 1024, MAX_FILE_BYTES + 1 - total))
            if not chunk:
                break
            chunks.append(chunk)
            total += len(chunk)
        after = os.fstat(descriptor)
    except FileHandlerError:
        raise
    except (OSError, ValueError):
        _raise_file_error("FILE_OPEN_FAILED")
    finally:
        os.close(descriptor)

    if after.st_ino != info.st_ino or after.st_dev != info.st_dev:
        _raise_file_error("FILE_CHANGED")
    if total > MAX_FILE_BYTES:
        _raise_file_error("FILE_TOO_LARGE")
    return b"".join(chunks)


@contextmanager
def _silence_worker_output():
    """Discard parser output inside this worker only."""

    previous_disable = logging.root.manager.disable
    saved_stdout = os.dup(1)
    saved_stderr = os.dup(2)
    null_descriptor = os.open(os.devnull, os.O_WRONLY)
    logging.disable(logging.CRITICAL)
    try:
        os.dup2(null_descriptor, 1)
        os.dup2(null_descriptor, 2)
        with redirect_stdout(io.StringIO()), redirect_stderr(io.StringIO()):
            yield
    finally:
        os.dup2(saved_stdout, 1)
        os.dup2(saved_stderr, 2)
        os.close(saved_stdout)
        os.close(saved_stderr)
        os.close(null_descriptor)
        logging.disable(previous_disable)


def _set_worker_limits() -> None:
    """Apply best-effort CPU and address-space limits in a worker."""

    if sys.platform == "win32":
        return
    try:
        import resource
    except ImportError:
        return

    for name, requested_soft in (
        ("RLIMIT_CPU", FILE_WORKER_CPU_SECONDS),
        ("RLIMIT_AS", FILE_WORKER_ADDRESS_SPACE_BYTES),
    ):
        limit = getattr(resource, name, None)
        if limit is None:
            continue
        try:
            current_soft, current_hard = resource.getrlimit(limit)
            selected_soft = (
                requested_soft
                if current_soft == resource.RLIM_INFINITY
                else min(current_soft, requested_soft)
            )
            if current_hard != resource.RLIM_INFINITY:
                selected_soft = min(selected_soft, current_hard)
            resource.setrlimit(limit, (selected_soft, current_hard))
        except (OSError, ValueError):
            continue


def _safe_payload(payload: Mapping[str, object]) -> bytes:
    """Encode a worker response with a strict size bound."""

    try:
        encoded = json.dumps(payload, ensure_ascii=False, separators=(",", ":")).encode("utf-8")
    except (TypeError, UnicodeError):
        encoded = b'{"ok":false,"code":"FILE_RESPONSE_INVALID"}'
    if len(encoded) > MAX_FILE_RESPONSE_BYTES:
        encoded = b'{"ok":false,"code":"FILE_RESPONSE_TOO_LARGE"}'
    return encoded


def _send_worker_result(connection: object, payload: Mapping[str, object]) -> None:
    try:
        connection.send_bytes(_safe_payload(payload))  # type: ignore[attr-defined]
    except (BrokenPipeError, EOFError, OSError):
        return


def _parse_plain_bytes(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        _raise_file_error("FILE_NOT_UTF8")


def _parse_excel_bytes(data: bytes) -> tuple[str, list[tuple[str, str]]]:
    _check_import("openpyxl", "openpyxl>=3.1.0")
    import openpyxl

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    except FileHandlerError:
        raise
    except Exception:
        _raise_file_error("FILE_MALFORMED")

    cells: list[tuple[str, str]] = []
    text_parts: list[str] = []
    try:
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            for row in worksheet.iter_rows():
                for cell in row:
                    if cell.value is not None:
                        cell_text = str(cell.value)
                        cells.append((f"{sheet_name}!{cell.coordinate}", cell_text))
                        text_parts.append(cell_text)
    except Exception:
        _raise_file_error("FILE_MALFORMED")
    finally:
        workbook.close()
    return "\n".join(text_parts), cells


def _parse_docx_bytes(data: bytes) -> tuple[str, list[tuple[str, str]]]:
    _check_import("docx", "python-docx>=1.1.0")
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except FileHandlerError:
        raise
    except Exception:
        _raise_file_error("FILE_MALFORMED")

    cells: list[tuple[str, str]] = []
    text_parts: list[str] = []
    try:
        for index, paragraph in enumerate(document.paragraphs):
            if paragraph.text.strip():
                cells.append((f"para:{index}", paragraph.text))
                text_parts.append(paragraph.text)
        for table_index, table in enumerate(document.tables):
            for row_index, row in enumerate(table.rows):
                for column_index, cell in enumerate(row.cells):
                    if cell.text.strip():
                        key = f"table:{table_index}:r{row_index}:c{column_index}"
                        cells.append((key, cell.text))
                        text_parts.append(cell.text)
    except Exception:
        _raise_file_error("FILE_MALFORMED")
    return "\n".join(text_parts), cells


def _parse_pdf_bytes(data: bytes) -> str:
    _check_import("pdfplumber", "pdfplumber>=0.11.0")
    import pdfplumber

    try:
        with pdfplumber.open(
            io.BytesIO(data), strict_metadata=False, raise_unicode_errors=True
        ) as pdf:
            if not 1 <= len(pdf.pages) <= PDF_MAX_PAGES:
                _raise_file_error("FILE_MALFORMED")
            page_texts: list[str] = []
            for page in pdf.pages:
                extracted = page.extract_text() or ""
                if not isinstance(extracted, str):
                    _raise_file_error("FILE_MALFORMED")
                page_texts.append(extracted)
    except FileHandlerError:
        raise
    except (ImportError, ModuleNotFoundError):
        _raise_file_error("FORMAT_UNAVAILABLE")
    except Exception:
        _raise_file_error("FILE_MALFORMED")

    text = unicodedata.normalize("NFKC", "\n\n".join(page_texts))
    if len(text.encode("utf-8")) > MAX_FILE_RESPONSE_BYTES:
        _raise_file_error("FILE_RESPONSE_TOO_LARGE")
    return text


def _parse_bytes(extension: str, data: bytes) -> dict[str, object]:
    if extension in PLAIN_TEXT_EXTENSIONS:
        return {"file_type": "plain", "text": _parse_plain_bytes(data), "cells": []}
    if extension in EXCEL_EXTENSIONS:
        text, cells = _parse_excel_bytes(data)
        return {"file_type": "excel", "text": text, "cells": cells}
    if extension in WORD_EXTENSIONS:
        text, cells = _parse_docx_bytes(data)
        return {"file_type": "docx", "text": text, "cells": cells}
    if extension in PDF_EXTENSIONS:
        return {"file_type": "pdf", "text": _parse_pdf_bytes(data), "cells": []}
    _raise_file_error("FILE_UNSUPPORTED")


def _anonymize_cell(cell_text: str, mapping: Mapping[str, str]) -> str:
    """Apply reverse mapping to a cell; mapping is placeholder -> original."""

    reverse = {value: key for key, value in mapping.items()}
    result = cell_text
    for original in sorted(reverse, key=len, reverse=True):
        result = result.replace(original, reverse[original])
    return result


def _write_excel_bytes(data: bytes, mapping: Mapping[str, str]) -> bytes:
    _check_import("openpyxl", "openpyxl>=3.1.0")
    import openpyxl

    try:
        workbook = openpyxl.load_workbook(io.BytesIO(data))
    except FileHandlerError:
        raise
    except Exception:
        _raise_file_error("FILE_MALFORMED")
    try:
        for sheet_name in workbook.sheetnames:
            worksheet = workbook[sheet_name]
            for row in worksheet.iter_rows():
                for cell in row:
                    if isinstance(cell.value, str):
                        cell.value = _anonymize_cell(cell.value, mapping)
        output = io.BytesIO()
        workbook.save(output)
        return output.getvalue()
    except Exception:
        _raise_file_error("FILE_WRITE_FAILED")
    finally:
        workbook.close()


def _replace_docx_paragraph(paragraph: object, anonymized: str) -> None:
    runs = getattr(paragraph, "runs", [])
    if runs:
        first_run = runs[0]
        for run in runs[1:]:
            run.text = ""
        first_run.text = anonymized
    else:
        setattr(paragraph, "text", anonymized)


def _write_docx_bytes(data: bytes, mapping: Mapping[str, str]) -> bytes:
    _check_import("docx", "python-docx>=1.1.0")
    from docx import Document

    try:
        document = Document(io.BytesIO(data))
    except FileHandlerError:
        raise
    except Exception:
        _raise_file_error("FILE_MALFORMED")
    try:
        for paragraph in document.paragraphs:
            if paragraph.text:
                _replace_docx_paragraph(paragraph, _anonymize_cell(paragraph.text, mapping))
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text and cell.paragraphs:
                        _replace_docx_paragraph(
                            cell.paragraphs[0], _anonymize_cell(cell.text, mapping)
                        )
        output = io.BytesIO()
        document.save(output)
        return output.getvalue()
    except Exception:
        _raise_file_error("FILE_WRITE_FAILED")


def _normalised_character_indices(
    chars: Sequence[Mapping[str, object]],
) -> tuple[str, list[int]]:
    """Build an NFKC stream and map each stream code point to a source char."""

    stream_parts: list[str] = []
    indices: list[int] = []
    for index, char in enumerate(chars):
        value = char.get("text")
        if not isinstance(value, str):
            _raise_file_error("PDF_OUTPUT_INVALID")
        normalised = unicodedata.normalize("NFKC", value)
        stream_parts.append(normalised)
        indices.extend([index] * len(normalised))
    return "".join(stream_parts), indices


def _pdf_value_matches(
    chars: Sequence[Mapping[str, object]], value: str
) -> list[list[Mapping[str, object]]]:
    """Return character groups for every occurrence of a value on a page."""

    if not value:
        _raise_file_error("PDF_REDACTION_UNRESOLVED")
    stream, indices = _normalised_character_indices(chars)
    needle = unicodedata.normalize("NFKC", value)
    matches: list[list[Mapping[str, object]]] = []
    cursor = 0
    while True:
        found = stream.find(needle, cursor)
        if found < 0:
            break
        end = found + len(needle)
        selected = sorted(set(indices[found:end]))
        matches.append([chars[index] for index in selected])
        cursor = end
    if matches:
        return matches

    # Some producers omit layout whitespace from character streams. Match a
    # compact form as a conservative fallback while retaining coordinates.
    compact_stream = "".join(char for char in stream if not char.isspace())
    compact_indices = [
        index for char, index in zip(stream, indices, strict=False) if not char.isspace()
    ]
    compact_needle = "".join(char for char in needle if not char.isspace())
    cursor = 0
    while compact_needle:
        found = compact_stream.find(compact_needle, cursor)
        if found < 0:
            break
        end = found + len(compact_needle)
        selected = sorted(set(compact_indices[found:end]))
        matches.append([chars[index] for index in selected])
        cursor = end
    return matches


def _char_box(char: Mapping[str, object]) -> tuple[float, float, float, float] | None:
    try:
        x0 = float(cast(str | int | float, char["x0"]))
        x1 = float(cast(str | int | float, char["x1"]))
        top = float(cast(str | int | float, char["top"]))
        bottom = float(cast(str | int | float, char["bottom"]))
    except (KeyError, TypeError, ValueError):
        return None
    values = (x0, x1, top, bottom)
    if any(not value == value or value in (float("inf"), float("-inf")) for value in values):
        return None
    return min(x0, x1), min(top, bottom), max(x0, x1), max(top, bottom)


def _group_boxes(
    chars: Sequence[Mapping[str, object]],
) -> list[tuple[float, float, float, float]]:
    boxes = [_char_box(char) for char in chars]
    valid = [box for box in boxes if box is not None]
    if not valid:
        return []
    lines: list[list[tuple[float, float, float, float]]] = []
    for box in sorted(valid, key=lambda item: (item[1], item[0])):
        for line in lines:
            line_top = min(item[1] for item in line)
            line_bottom = max(item[3] for item in line)
            if box[1] <= line_bottom and box[3] >= line_top:
                line.append(box)
                break
        else:
            lines.append([box])
    return [
        (
            min(item[0] for item in line),
            min(item[1] for item in line),
            max(item[2] for item in line),
            max(item[3] for item in line),
        )
        for line in lines
    ]


def _write_pdf_bytes(data: bytes, mapping: Mapping[str, str]) -> bytes:
    """Rasterise every PDF page after masking every mapped source value."""

    try:
        import pdfplumber
        import pypdfium2 as pdfium
        from PIL import ImageDraw, ImageFont
    except (ImportError, ModuleNotFoundError):
        _raise_file_error("PDF_RENDER_UNAVAILABLE")

    try:
        with pdfplumber.open(
            io.BytesIO(data), strict_metadata=False, raise_unicode_errors=True
        ) as parsed_pdf:
            if not 1 <= len(parsed_pdf.pages) <= PDF_MAX_PAGES:
                _raise_file_error("PDF_OUTPUT_INVALID")
            try:
                rendered_document = pdfium.PdfDocument(data)
            except Exception:
                _raise_file_error("PDF_OUTPUT_INVALID")

            images: list[object] = []
            try:
                if len(rendered_document) != len(parsed_pdf.pages):
                    _raise_file_error("PDF_OUTPUT_INVALID")
                font = ImageFont.load_default()
                mapping_items: list[tuple[str, str]] = []
                for placeholder, original in mapping.items():
                    if (
                        not isinstance(placeholder, str)
                        or not isinstance(original, str)
                        or not original
                    ):
                        _raise_file_error("PDF_REDACTION_UNRESOLVED")
                    mapping_items.append((placeholder, original))
                found_values: set[str] = set()
                for page_index, parsed_page in enumerate(parsed_pdf.pages):
                    try:
                        rendered_page = rendered_document[page_index]
                        rendered_width = float(rendered_page.get_width())
                        rendered_height = float(rendered_page.get_height())
                        bitmap = rendered_page.render(scale=PDF_RENDER_SCALE)
                        image = bitmap.to_pil().convert("RGB")
                    except Exception:
                        _raise_file_error("PDF_OUTPUT_INVALID")

                    # pdfplumber's char coordinates are relative to the
                    # page's CropBox (not necessarily the MediaBox origin),
                    # and that is also what pypdfium2 renders -- see
                    # pdfplumber.display.PageImage._reproject. Scale from
                    # the CropBox, not the raw rendered pixel size, so a
                    # non-zero MediaBox origin or a CropBox smaller than the
                    # MediaBox still lines the mask up with the glyphs.
                    try:
                        crop_x0, crop_top, crop_x1, crop_bottom = (
                            float(value) for value in parsed_page.cropbox
                        )
                    except Exception:
                        _raise_file_error("PDF_OUTPUT_INVALID")
                    if any(
                        not value == value or value in (float("inf"), float("-inf"))
                        for value in (crop_x0, crop_top, crop_x1, crop_bottom)
                    ):
                        _raise_file_error("PDF_OUTPUT_INVALID")
                    crop_width = crop_x1 - crop_x0
                    crop_height = crop_bottom - crop_top
                    if crop_width <= 0 or crop_height <= 0:
                        _raise_file_error("PDF_OUTPUT_INVALID")
                    # Fail closed if pypdfium2 did not actually render the
                    # CropBox extent we are about to scale coordinates
                    # against (e.g. it clipped to the MediaBox instead).
                    if (
                        abs(rendered_width - crop_width) > PDF_CROPBOX_RENDER_TOLERANCE_POINTS
                        or abs(rendered_height - crop_height) > PDF_CROPBOX_RENDER_TOLERANCE_POINTS
                    ):
                        _raise_file_error("PDF_OUTPUT_INVALID")

                    chars = [char for char in parsed_page.chars if isinstance(char, dict)]
                    for placeholder, original in mapping_items:
                        matches = _pdf_value_matches(chars, original)
                        if matches:
                            found_values.add(original)
                        for match in matches:
                            for x0, top, x1, bottom in _group_boxes(match):
                                x_scale = image.width / crop_width
                                y_scale = image.height / crop_height
                                padding_x = max(2.0, x_scale * 1.0)
                                padding_y = max(2.0, y_scale * 1.0)
                                left = max(0.0, (x0 - crop_x0) * x_scale - padding_x)
                                upper = max(0.0, (top - crop_top) * y_scale - padding_y)
                                right = min(
                                    float(image.width), (x1 - crop_x0) * x_scale + padding_x
                                )
                                lower = min(
                                    float(image.height),
                                    (bottom - crop_top) * y_scale + padding_y,
                                )
                                draw = ImageDraw.Draw(image)
                                draw.rectangle((left, upper, right, lower), fill="white")
                                draw.rectangle((left, upper, right, lower), outline="black")
                                draw.text(
                                    (left + padding_x, upper + padding_y),
                                    placeholder,
                                    fill="black",
                                    font=font,
                                )
                    images.append(image)

                if not images:
                    _raise_file_error("PDF_OUTPUT_INVALID")
                if any(original not in found_values for _, original in mapping_items):
                    _raise_file_error("PDF_REDACTION_UNRESOLVED")
                output = io.BytesIO()
                first = images[0]
                first.save(  # type: ignore[attr-defined]
                    output,
                    format="PDF",
                    save_all=True,
                    append_images=images[1:],
                    resolution=72.0 * PDF_RENDER_SCALE,
                )
                generated = output.getvalue()
            finally:
                try:
                    rendered_document.close()
                except Exception:
                    pass
    except FileHandlerError:
        raise
    except (ImportError, ModuleNotFoundError):
        _raise_file_error("PDF_RENDER_UNAVAILABLE")
    except Exception:
        _raise_file_error("PDF_OUTPUT_INVALID")

    if not generated or len(generated) > MAX_FILE_RESPONSE_BYTES:
        _raise_file_error("PDF_OUTPUT_INVALID")

    # The output is image-only. Reopen it in the worker to catch accidental
    # text-layer/metadata paths introduced by a future renderer change.
    try:
        with pdfplumber.open(io.BytesIO(generated), strict_metadata=False) as check_pdf:
            if len(check_pdf.pages) != len(images):
                _raise_file_error("PDF_OUTPUT_INVALID")
            extracted = "\n".join(page.extract_text() or "" for page in check_pdf.pages)
            if any(value and value in extracted for value in mapping.values()):
                _raise_file_error("PDF_OUTPUT_INVALID")
    except FileHandlerError:
        raise
    except Exception:
        _raise_file_error("PDF_OUTPUT_INVALID")
    return generated


def _write_bytes_for_operation(extension: str, data: bytes, mapping: Mapping[str, str]) -> bytes:
    if extension in EXCEL_EXTENSIONS:
        return _write_excel_bytes(data, mapping)
    if extension in WORD_EXTENSIONS:
        return _write_docx_bytes(data, mapping)
    if extension in PDF_EXTENSIONS:
        return _write_pdf_bytes(data, mapping)
    _raise_file_error("FILE_UNSUPPORTED")


def _file_worker(
    input_connection: object,
    output_connection: object,
    operation: str,
    extension: str,
    mapping: Mapping[str, str],
) -> None:
    """Spawn-safe target for all third-party format parsing and writing."""

    with _silence_worker_output():
        try:
            _set_worker_limits()
            data = input_connection.recv_bytes(MAX_FILE_BYTES + 1)  # type: ignore[attr-defined]
            if len(data) > MAX_FILE_BYTES:
                _send_worker_result(output_connection, {"ok": False, "code": "FILE_TOO_LARGE"})
            elif operation == "read":
                payload = _parse_bytes(extension, data)
                payload["ok"] = True
                _send_worker_result(output_connection, payload)
            elif operation == "write":
                output = _write_bytes_for_operation(extension, data, mapping)
                if len(output) > MAX_FILE_RESPONSE_BYTES:
                    _send_worker_result(
                        output_connection, {"ok": False, "code": "FILE_RESPONSE_TOO_LARGE"}
                    )
                else:
                    _send_worker_result(
                        output_connection,
                        {"ok": True, "data": base64.b64encode(output).decode("ascii")},
                    )
            else:
                _send_worker_result(
                    output_connection, {"ok": False, "code": "FILE_RESPONSE_INVALID"}
                )
        except FileHandlerError as error:
            _send_worker_result(output_connection, {"ok": False, "code": error.code})
        except MemoryError:
            _send_worker_result(
                output_connection, {"ok": False, "code": "FILE_PARSE_RESOURCE_LIMIT"}
            )
        except BaseException:
            _send_worker_result(output_connection, {"ok": False, "code": "FILE_MALFORMED"})
        finally:
            for connection in (input_connection, output_connection):
                try:
                    connection.close()  # type: ignore[attr-defined]
                except (AttributeError, OSError):
                    pass


def _stop_worker(process: BaseProcess) -> None:
    """Terminate then kill a child and always reap it."""

    try:
        alive = process.is_alive()
    except AssertionError:
        return
    if alive:
        process.terminate()
        process.join(FILE_PARSE_GRACE_SECONDS)
    try:
        alive = process.is_alive()
    except AssertionError:
        return
    if alive and hasattr(process, "kill"):
        process.kill()
        process.join(FILE_PARSE_GRACE_SECONDS)
    elif not alive:
        process.join(FILE_PARSE_GRACE_SECONDS)


def _run_file_worker(
    operation: str,
    extension: str,
    data: bytes,
    mapping: Mapping[str, str] | None = None,
) -> dict[str, object]:
    """Run a worker with bounded IPC and a wall-clock deadline."""

    if not isinstance(data, bytes):
        _raise_file_error("FILE_RESPONSE_INVALID")
    if len(data) > MAX_FILE_BYTES:
        _raise_file_error("FILE_TOO_LARGE")
    context = multiprocessing.get_context("spawn")
    input_read, input_write = context.Pipe(duplex=False)
    output_read, output_write = context.Pipe(duplex=False)
    process = context.Process(
        target=_file_worker,
        args=(input_read, output_write, operation, extension, dict(mapping or {})),
        daemon=True,
    )
    sender_errors: list[BaseException] = []

    def send_input() -> None:
        try:
            input_write.send_bytes(data)
        except BaseException as error:
            sender_errors.append(error)

    sender: threading.Thread | None = None
    process_started = False
    payload: bytes | None = None
    try:
        process.start()
        process_started = True
        input_read.close()
        output_write.close()
        sender = threading.Thread(target=send_input, daemon=True)
        sender.start()
        deadline = time.monotonic() + FILE_PARSE_TIMEOUT_SECONDS
        while True:
            if output_read.poll(0.05):
                try:
                    payload = output_read.recv_bytes(MAX_FILE_RESPONSE_BYTES)
                except BufferTooShort:
                    _raise_file_error("FILE_RESPONSE_TOO_LARGE")
                except (EOFError, OSError):
                    _raise_file_error("FILE_IPC_FAILED")
                break
            if sender is not None and not sender.is_alive() and sender_errors:
                _raise_file_error("FILE_IPC_FAILED")
            if time.monotonic() >= deadline:
                _raise_file_error("FILE_PARSE_TIMEOUT")
            if not process.is_alive() and not output_read.poll(0):
                break

        if payload is None:
            if process.exitcode is not None and process.exitcode < 0:
                _raise_file_error("FILE_PARSE_RESOURCE_LIMIT")
            _raise_file_error("FILE_PARSE_CRASHED")
        if len(payload) > MAX_FILE_RESPONSE_BYTES:
            _raise_file_error("FILE_RESPONSE_TOO_LARGE")
        try:
            decoded = json.loads(payload.decode("utf-8"))
        except (UnicodeDecodeError, json.JSONDecodeError):
            _raise_file_error("FILE_RESPONSE_INVALID")
        if not isinstance(decoded, dict) or not isinstance(decoded.get("ok"), bool):
            _raise_file_error("FILE_RESPONSE_INVALID")
        if not decoded["ok"]:
            code = decoded.get("code")
            _raise_file_error(code if isinstance(code, str) else "FILE_RESPONSE_INVALID")
        process.join(FILE_PARSE_GRACE_SECONDS)
        return decoded
    except FileHandlerError:
        raise
    except (OSError, ValueError):
        _raise_file_error("FILE_IPC_FAILED")
    finally:
        if sender is not None:
            sender.join(FILE_PARSE_GRACE_SECONDS)
        if process_started:
            _stop_worker(process)
        for connection in (input_write, output_read):
            try:
                connection.close()
            except (AttributeError, OSError):
                pass


def _decoded_cells(value: object) -> list[tuple[str, str]]:
    if not isinstance(value, list):
        _raise_file_error("FILE_RESPONSE_INVALID")
    cells: list[tuple[str, str]] = []
    for item in value:
        if not isinstance(item, list) or len(item) != 2:
            _raise_file_error("FILE_RESPONSE_INVALID")
        key, text = item
        if not isinstance(key, str) or not isinstance(text, str):
            _raise_file_error("FILE_RESPONSE_INVALID")
        cells.append((key, text))
    return cells


def read_file(path: str | Path) -> FileContent:
    """Read a file through an isolated, bounded parser worker."""

    try:
        file_path = Path(path)
        extension = file_path.suffix.lower()
    except (TypeError, ValueError):
        _raise_file_error("FILE_UNSUPPORTED")
    if extension not in SUPPORTED_EXTENSIONS:
        _raise_file_error("FILE_UNSUPPORTED")
    source_bytes = _read_source_bytes(file_path)
    payload = _run_file_worker("read", extension, source_bytes)
    file_type = payload.get("file_type")
    text = payload.get("text")
    if file_type not in {"plain", "excel", "docx", "pdf"} or not isinstance(text, str):
        _raise_file_error("FILE_RESPONSE_INVALID")
    return FileContent(
        text=text,
        file_type=file_type,
        source_path=str(file_path),
        cells=_decoded_cells(payload.get("cells")),
        source_bytes=source_bytes,
    )


# Compatibility helpers for callers that imported the old private functions.
# Public read_file never invokes these in the parent process.
def _read_plain(path: Path) -> FileContent:
    data = _read_source_bytes(path)
    return FileContent(
        text=_parse_plain_bytes(data),
        file_type="plain",
        source_path=str(path),
        source_bytes=data,
    )


def _read_excel(path: Path) -> FileContent:
    """Compatibility wrapper that keeps Excel parsing in the worker."""

    return read_file(path)


def _read_docx(path: Path) -> FileContent:
    """Compatibility wrapper that keeps Word parsing in the worker."""

    return read_file(path)


def _read_pdf(path: Path) -> FileContent:
    """Compatibility wrapper that keeps PDF parsing in the worker."""

    return read_file(path)


def _source_bytes(content: FileContent) -> bytes:
    if content.source_bytes:
        return content.source_bytes
    if content.source_path:
        return _read_source_bytes(Path(content.source_path))
    _raise_file_error("FILE_OPEN_FAILED")


def _write_output_bytes(output_path: Path, data: bytes) -> None:
    """Write worker output without following a symlink."""

    try:
        output = output_path.expanduser()
        if output.exists() or output.is_symlink():
            if output.is_symlink() or output.is_dir():
                _raise_file_error("FILE_WRITE_FAILED")
        output.parent.mkdir(parents=True, exist_ok=True)
        if output.parent.is_symlink():
            _raise_file_error("FILE_WRITE_FAILED")
        flags = os.O_WRONLY | os.O_CREAT | os.O_TRUNC | getattr(os, "O_NOFOLLOW", 0)
        descriptor = os.open(output, flags, 0o600)
    except FileHandlerError:
        raise
    except (OSError, ValueError):
        _raise_file_error("FILE_WRITE_FAILED")
    try:
        with os.fdopen(descriptor, "wb") as handle:
            descriptor = -1
            handle.write(data)
            handle.flush()
            os.fsync(handle.fileno())
    except (OSError, ValueError):
        if descriptor >= 0:
            os.close(descriptor)
        _raise_file_error("FILE_WRITE_FAILED")


def _extension_for_type(file_type: str) -> str:
    if file_type == "excel":
        return ".xlsx"
    if file_type == "docx":
        return ".docx"
    if file_type == "pdf":
        return ".pdf"
    _raise_file_error("FILE_UNSUPPORTED")


def write_file(
    content: FileContent,
    anonymized_text: str,
    mapping: dict[str, str],
    output_path: str | Path,
) -> None:
    """Write anonymized content back to a file, preserving its format.

    Structured formats and PDF output are processed in a fresh worker. PDF
    output is a layout-sized raster PDF; unresolved redactions fail closed and
    leave no output artifact.
    """

    out = Path(output_path)
    if content.file_type == "plain":
        try:
            out.write_text(anonymized_text, encoding="utf-8")
        except (OSError, ValueError):
            _raise_file_error("FILE_WRITE_FAILED")
        return
    extension = _extension_for_type(content.file_type)
    payload = _run_file_worker("write", extension, _source_bytes(content), mapping)
    encoded = payload.get("data")
    if not isinstance(encoded, str):
        _raise_file_error("FILE_RESPONSE_INVALID")
    try:
        generated = base64.b64decode(encoded.encode("ascii"), validate=True)
    except (ValueError, UnicodeEncodeError):
        _raise_file_error("FILE_RESPONSE_INVALID")
    if not generated or len(generated) > MAX_FILE_RESPONSE_BYTES:
        _raise_file_error("FILE_RESPONSE_INVALID")
    _write_output_bytes(out, generated)


def _write_excel(content: FileContent, mapping: dict[str, str], out: Path) -> None:
    """Compatibility wrapper that keeps Excel writing in the worker."""

    write_file(content, content.text, mapping, out)


def _write_docx(content: FileContent, mapping: dict[str, str], out: Path) -> None:
    """Compatibility wrapper that keeps Word writing in the worker."""

    write_file(content, content.text, mapping, out)


def is_supported(path: str | Path) -> bool:
    """Return True if the file extension is supported."""

    try:
        return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS
    except (TypeError, ValueError):
        return False


def get_output_extension(input_path: str | Path) -> str:
    """Return the output extension for a supported input file."""

    ext = Path(input_path).suffix.lower()
    return ".pdf" if ext in PDF_EXTENSIONS else ext
