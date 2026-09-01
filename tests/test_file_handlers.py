"""Tests for multi-format file reading and writing (file_handlers.py)."""

from __future__ import annotations

from pathlib import Path

import pytest

from pii_guard.file_handlers import (
    FileHandlerError,
    _read_source_bytes,
    _write_output_bytes,
    get_output_extension,
    is_supported,
    read_file,
    write_file,
)


def _two_page_pdf_bytes() -> bytes:
    """Build a dependency-free two-page PDF with one PII value per page."""

    def stream(body: bytes) -> bytes:
        return (
            b"<< /Length "
            + str(len(body)).encode("ascii")
            + b" >>\nstream\n"
            + body
            + b"\nendstream"
        )

    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R 4 0 R] /Count 2 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 6 0 R >>"
        ),
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 7 0 R >>"
        ),
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        stream(b"BT /F1 18 Tf 100 700 Td (ID: A123456789) Tj ET"),
        stream(b"BT /F1 18 Tf 100 700 Td (Phone: 0912345678) Tj ET"),
    ]
    document = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, value in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{object_number} 0 obj\n".encode("ascii"))
        document.extend(value)
        document.extend(b"\nendobj\n")
    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(document)


def _boxed_pdf_bytes(
    media_box: str,
    extra_page_entries: str = "",
    text: str = "SECRET A123456789",
    position: tuple[float, float] = (120, 600),
    font_size: int = 24,
) -> bytes:
    """Build a dependency-free one-page PDF with a configurable page box.

    ``extra_page_entries`` lets a case add e.g. ``/CropBox [...]`` or
    ``/Rotate 90`` to the page dictionary -- the two attributes that make a
    page's CropBox differ from its MediaBox, which is what the PDF
    redaction mask-alignment fix in ``_write_pdf_bytes`` has to handle.
    """

    def stream(body: bytes) -> bytes:
        return (
            b"<< /Length "
            + str(len(body)).encode("ascii")
            + b" >>\nstream\n"
            + body
            + b"\nendstream"
        )

    x, y = position
    content = f"BT /F1 {font_size} Tf {x:g} {y:g} Td ({text}) Tj ET".encode("ascii")
    page_dict = (
        f"<< /Type /Page /Parent 2 0 R /MediaBox {media_box} {extra_page_entries} "
        "/Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>"
    ).encode("ascii")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        page_dict,
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        stream(content),
    ]
    document = bytearray(b"%PDF-1.4\n%\xe2\xe3\xcf\xd3\n")
    offsets = [0]
    for object_number, value in enumerate(objects, start=1):
        offsets.append(len(document))
        document.extend(f"{object_number} 0 obj\n".encode("ascii"))
        document.extend(value)
        document.extend(b"\nendobj\n")
    xref_offset = len(document)
    document.extend(f"xref\n0 {len(objects) + 1}\n".encode("ascii"))
    document.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        document.extend(f"{offset:010d} 00000 n \n".encode("ascii"))
    document.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode("ascii")
    )
    return bytes(document)


def _render_pdf_page(data: bytes, scale: float = 1.0):
    """Render a single-page PDF's first page to a PIL RGB image."""

    import pypdfium2 as pdfium

    document = pdfium.PdfDocument(data)
    try:
        return document[0].render(scale=scale).to_pil().convert("RGB")
    finally:
        document.close()


def _dark_pixel_mask(image) -> bytes:
    """Return a 0/1-per-pixel dark mask (grayscale < 128) as raw bytes."""

    return image.convert("L").point(lambda level: 1 if level < 128 else 0).tobytes()


def _still_dark_fraction(source_image, output_image) -> float:
    """Fraction of originally-dark source pixels that are still dark in output."""

    source_mask = _dark_pixel_mask(source_image)
    output_mask = _dark_pixel_mask(output_image)
    assert len(source_mask) == len(output_mask)
    source_dark = source_mask.count(1)
    if source_dark == 0:
        return 0.0
    still_dark = sum(1 for s, o in zip(source_mask, output_mask, strict=True) if s and o)
    return still_dark / source_dark


# ---------------------------------------------------------------------------
# is_supported / get_output_extension
# ---------------------------------------------------------------------------


class TestHelpers:
    def test_supported_txt(self):
        assert is_supported("data.txt")

    def test_supported_csv(self):
        assert is_supported("data.csv")

    def test_supported_xlsx(self):
        assert is_supported("report.xlsx")

    def test_supported_docx(self):
        assert is_supported("memo.docx")

    def test_supported_pdf(self):
        assert is_supported("invoice.pdf")

    def test_unsupported_pptx(self):
        assert not is_supported("slides.pptx")

    def test_output_ext_txt(self):
        assert get_output_extension("data.txt") == ".txt"

    def test_output_ext_xlsx(self):
        assert get_output_extension("report.xlsx") == ".xlsx"

    def test_output_ext_pdf_is_pdf(self):
        assert get_output_extension("invoice.pdf") == ".pdf"


# ---------------------------------------------------------------------------
# Plain text reading and writing
# ---------------------------------------------------------------------------


class TestPlainText:
    def test_read_txt(self, tmp_path: Path):
        f = tmp_path / "data.txt"
        f.write_text("身分證A123456789", encoding="utf-8")
        content = read_file(f)
        assert content.file_type == "plain"
        assert "A123456789" in content.text

    def test_read_csv(self, tmp_path: Path):
        f = tmp_path / "data.csv"
        f.write_text("name,id\n張三,A123456789", encoding="utf-8")
        content = read_file(f)
        assert content.file_type == "plain"
        assert "A123456789" in content.text

    def test_write_plain(self, tmp_path: Path):
        f = tmp_path / "data.txt"
        f.write_text("身分證A123456789", encoding="utf-8")
        content = read_file(f)
        out = tmp_path / "out.txt"
        write_file(content, "身分證<TW_NATIONAL_ID_1>", {}, out)
        assert "<TW_NATIONAL_ID_1>" in out.read_text(encoding="utf-8")

    def test_file_not_found(self):
        with pytest.raises(FileHandlerError, match="FILE_NOT_FOUND"):
            read_file("/tmp/definitely_not_exist_pii_guard.txt")

    def test_unsupported_extension(self, tmp_path: Path):
        f = tmp_path / "data.pptx"
        f.write_bytes(b"fake")
        with pytest.raises(FileHandlerError, match="FILE_UNSUPPORTED"):
            read_file(f)


# ---------------------------------------------------------------------------
# Excel (.xlsx) reading and writing
# ---------------------------------------------------------------------------


class TestExcel:
    @pytest.fixture()
    def sample_xlsx(self, tmp_path: Path) -> Path:
        import openpyxl

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = "Sheet1"
        ws["A1"] = "姓名"
        ws["B1"] = "身分證"
        ws["A2"] = "張三"
        ws["B2"] = "A123456789"
        ws["A3"] = "李四"
        ws["B3"] = "B234567890"
        path = tmp_path / "data.xlsx"
        wb.save(str(path))
        wb.close()
        return path

    def test_read_xlsx(self, sample_xlsx: Path):
        content = read_file(sample_xlsx)
        assert content.file_type == "excel"
        assert "A123456789" in content.text
        assert "B234567890" in content.text
        assert len(content.cells) == 6  # 3 rows * 2 cols

    def test_read_xlsx_cells_have_keys(self, sample_xlsx: Path):
        content = read_file(sample_xlsx)
        keys = [k for k, _ in content.cells]
        assert any("Sheet1!" in k for k in keys)

    def test_write_xlsx_roundtrip(self, sample_xlsx: Path, tmp_path: Path, spacy_only_engine):
        content = read_file(sample_xlsx)
        anonymized_text, mapping = spacy_only_engine.anonymize(content.text)

        out = tmp_path / "anon.xlsx"
        write_file(content, anonymized_text, mapping, out)

        # Read back and verify PII is replaced
        import openpyxl

        wb = openpyxl.load_workbook(str(out), read_only=True)
        ws = wb.active
        values = [str(c.value) for row in ws.iter_rows() for c in row if c.value]
        wb.close()
        assert "A123456789" not in " ".join(values)
        assert "B234567890" not in " ".join(values)


# ---------------------------------------------------------------------------
# Word (.docx) reading and writing
# ---------------------------------------------------------------------------


class TestDocx:
    @pytest.fixture()
    def sample_docx(self, tmp_path: Path) -> Path:
        from docx import Document

        doc = Document()
        doc.add_paragraph("客戶：張三")
        doc.add_paragraph("身分證字號：A123456789")
        doc.add_paragraph("手機：0912345678")

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "姓名"
        table.rows[0].cells[1].text = "Email"
        table.rows[1].cells[0].text = "李四"
        table.rows[1].cells[1].text = "test@example.com"

        path = tmp_path / "data.docx"
        doc.save(str(path))
        return path

    def test_read_docx(self, sample_docx: Path):
        content = read_file(sample_docx)
        assert content.file_type == "docx"
        assert "A123456789" in content.text
        assert "0912345678" in content.text
        assert "test@example.com" in content.text

    def test_read_docx_has_table_cells(self, sample_docx: Path):
        content = read_file(sample_docx)
        keys = [k for k, _ in content.cells]
        assert any(k.startswith("table:") for k in keys)

    def test_write_docx_roundtrip(self, sample_docx: Path, tmp_path: Path, spacy_only_engine):
        content = read_file(sample_docx)
        anonymized_text, mapping = spacy_only_engine.anonymize(content.text)

        out = tmp_path / "anon.docx"
        write_file(content, anonymized_text, mapping, out)

        # Read back and verify PII is replaced
        from docx import Document

        doc = Document(str(out))
        all_text = " ".join(p.text for p in doc.paragraphs)
        assert "A123456789" not in all_text
        assert "0912345678" not in all_text


# ---------------------------------------------------------------------------
# PDF (.pdf) reading
# ---------------------------------------------------------------------------


class TestPdf:
    @pytest.fixture()
    def sample_pdf(self, tmp_path: Path) -> Path:
        """Create a minimal PDF with PII text using pdfplumber-compatible format."""
        # Use reportlab if available, otherwise create a minimal PDF manually
        try:
            from reportlab.pdfgen import canvas

            path = tmp_path / "data.pdf"
            c = canvas.Canvas(str(path))
            c.drawString(100, 750, "ID: A123456789")
            c.drawString(100, 730, "Phone: 0912345678")
            c.save()
            return path
        except ImportError:
            # Create a minimal valid PDF with text
            path = tmp_path / "data.pdf"
            pdf_content = (
                b"%PDF-1.4\n"
                b"1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n"
                b"2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n"
                b"3 0 obj<</Type/Page/MediaBox[0 0 612 792]"
                b"/Parent 2 0 R/Contents 4 0 R/Resources<</Font<</F1 5 0 R>>>>>>endobj\n"
                b"4 0 obj<</Length 44>>stream\n"
                b"BT /F1 12 Tf 100 750 Td (A123456789) Tj ET\n"
                b"endstream\nendobj\n"
                b"5 0 obj<</Type/Font/Subtype/Type1/BaseFont/Helvetica>>endobj\n"
                b"xref\n0 6\n"
                b"0000000000 65535 f \n"
                b"0000000009 00000 n \n"
                b"0000000058 00000 n \n"
                b"0000000115 00000 n \n"
                b"0000000266 00000 n \n"
                b"0000000360 00000 n \n"
                b"trailer<</Size 6/Root 1 0 R>>\n"
                b"startxref\n431\n%%EOF"
            )
            path.write_bytes(pdf_content)
            return path

    def test_read_pdf(self, sample_pdf: Path):
        content = read_file(sample_pdf)
        assert content.file_type == "pdf"
        assert "A123456789" in content.text

    def test_pdf_no_cells(self, sample_pdf: Path):
        content = read_file(sample_pdf)
        assert content.cells == []

    def test_write_pdf_preserves_pdf_layout_and_removes_text(
        self, sample_pdf: Path, tmp_path: Path
    ):
        content = read_file(sample_pdf)
        out = tmp_path / "output.pdf"
        import pypdfium2 as pdfium

        original_pdf = pdfium.PdfDocument(content.source_bytes)
        original_size = (original_pdf[0].get_width(), original_pdf[0].get_height())
        original_pdf.close()
        write_file(content, "", {"<TW_NATIONAL_ID_1>": "A123456789"}, out)
        assert out.exists()
        assert out.read_bytes().startswith(b"%PDF-")

        # The generated PDF is intentionally image-only.  Its page count and
        # physical size remain the same while the source PII is not exposed as
        # selectable text or embedded metadata.
        import pdfplumber

        with pdfplumber.open(str(out)) as generated:
            assert len(generated.pages) == 1
            assert "A123456789" not in "\n".join(
                page.extract_text() or "" for page in generated.pages
            )
        generated_pdf = pdfium.PdfDocument(out.read_bytes())
        assert (generated_pdf[0].get_width(), generated_pdf[0].get_height()) == pytest.approx(
            original_size, abs=0.5
        )
        generated_pdf.close()

    def test_pdf_write_fails_closed_if_mapping_cannot_be_located(
        self, sample_pdf: Path, tmp_path: Path
    ):
        content = read_file(sample_pdf)
        out = tmp_path / "must-not-exist.pdf"
        with pytest.raises(FileHandlerError) as error:
            write_file(content, "", {"<PII_1>": "not-present-in-document"}, out)
        assert error.value.code == "PDF_REDACTION_UNRESOLVED"
        assert not out.exists()

    def test_pdf_write_masks_values_across_pages(self, tmp_path: Path):
        source = tmp_path / "two-page.pdf"
        source.write_bytes(_two_page_pdf_bytes())
        content = read_file(source)
        out = tmp_path / "two-page-redacted.pdf"
        mapping = {
            "<TW_NATIONAL_ID_1>": "A123456789",
            "<TW_MOBILE_1>": "0912345678",
        }

        write_file(content, "", mapping, out)

        import pdfplumber

        with pdfplumber.open(str(out)) as generated:
            assert len(generated.pages) == 2
            extracted = "\n".join(page.extract_text() or "" for page in generated.pages)
        assert "A123456789" not in extracted
        assert "0912345678" not in extracted

    # -----------------------------------------------------------------
    # Mask/glyph coordinate alignment across page-box variants.
    #
    # ``_write_pdf_bytes`` used to scale pdfplumber's character coordinates
    # by the raw rendered pixel size (``rendered_page.get_width()/height()``
    # against the pixel size at ``PDF_RENDER_SCALE``). That only matches
    # pdfplumber's coordinate space when a page's CropBox equals its
    # MediaBox and both start at the origin. A non-zero MediaBox origin, or
    # a CropBox smaller than the MediaBox, drew the mask tens of points
    # away from the actual glyphs while still reporting success.
    # -----------------------------------------------------------------

    _PDF_BOX_CASES: dict[str, tuple[str, str]] = {
        "plain page": ("[0 0 612 792]", ""),
        "non-zero MediaBox origin": ("[100 100 712 892]", ""),
        "CropBox smaller than MediaBox": ("[0 0 612 792]", "/CropBox [50 400 562 750]"),
        "Rotate 90": ("[0 0 612 792]", "/Rotate 90"),
    }

    @pytest.mark.parametrize("case_name", list(_PDF_BOX_CASES))
    def test_mask_covers_glyphs_regardless_of_page_box(self, tmp_path: Path, case_name: str):
        pytest.importorskip("pdfplumber")
        pytest.importorskip("pypdfium2")
        pytest.importorskip("PIL")

        media_box, extra_page_entries = self._PDF_BOX_CASES[case_name]
        source = tmp_path / "source.pdf"
        source.write_bytes(_boxed_pdf_bytes(media_box, extra_page_entries))

        content = read_file(source)
        assert content.file_type == "pdf"
        assert "A123456789" in content.text

        out = tmp_path / "redacted.pdf"
        write_file(content, "", {"<TW_NATIONAL_ID_1>": "A123456789"}, out)
        assert out.exists()

        source_image = _render_pdf_page(content.source_bytes)
        output_image = _render_pdf_page(out.read_bytes())
        assert source_image.size == output_image.size

        # The fixture draws "SECRET A123456789"; only "A123456789" is
        # mapped, so "SECRET " (with its trailing space) is left un-redacted
        # by design. Empirically, roughly 58% of the rendered ink belongs to
        # "A123456789" and ~42% to the literal "SECRET " prefix, so a
        # correctly-aligned mask leaves about 42% of the source's dark
        # pixels still dark. The coordinate bug drew the mask tens of
        # points away from the glyphs, leaving ~99% still dark. 0.60 sits
        # comfortably above the correct value and below the bug's value.
        still_dark_fraction = _still_dark_fraction(source_image, output_image)
        assert still_dark_fraction < 0.60, (
            f"{case_name}: {still_dark_fraction:.2%} of the source's dark "
            "pixels are still dark after redaction -- the mask is "
            "misaligned with the glyphs"
        )

    # Byte-exactness of the raw file IO. Windows opens os.open descriptors in
    # text mode unless O_BINARY is set: the read then stops at the first 0x1A
    # and CRLF is rewritten, which silently truncated every xlsx/docx to a few
    # dozen bytes. These assertions are trivially true on POSIX and are the
    # regression guard on Windows.

    def test_read_source_bytes_is_byte_exact(self, tmp_path: Path):
        path = tmp_path / "binary.xlsx"
        payload = b"PK\x03\x04" + bytes(range(256)) + b"\r\n\x1a" + b"tail" * 64
        path.write_bytes(payload)

        assert _read_source_bytes(path) == payload

    def test_write_output_bytes_is_byte_exact(self, tmp_path: Path):
        out = tmp_path / "binary-out.xlsx"
        payload = b"PK\x03\x04" + bytes(range(256)) + b"\r\n\x1a" + b"tail" * 64
        _write_output_bytes(out, payload)

        assert out.read_bytes() == payload

    def test_file_errors_do_not_include_source_path(self, tmp_path: Path):
        private_name = "private-customer-name-A123456789"
        missing = tmp_path / private_name / "record.docx"
        with pytest.raises(FileHandlerError) as error:
            read_file(missing)
        rendered = str(error.value)
        assert private_name not in rendered
        assert str(tmp_path) not in rendered

    @pytest.mark.parametrize("extension", [".docx", ".xlsx", ".pdf"])
    def test_malformed_format_error_is_fixed_and_does_not_echo_input(
        self, tmp_path: Path, extension: str
    ):
        private_name = "private-customer-name-A123456789"
        path = tmp_path / f"{private_name}{extension}"
        path.write_bytes(b"malformed source A123456789")
        with pytest.raises(FileHandlerError) as error:
            read_file(path)
        rendered = str(error.value)
        assert error.value.code in {"FILE_MALFORMED", "FORMAT_UNAVAILABLE"}
        assert private_name not in rendered
        assert "A123456789" not in rendered


# ---------------------------------------------------------------------------
# Integration: full anonymize + write-back + restore roundtrip
# ---------------------------------------------------------------------------


class TestFullRoundtrip:
    def test_txt_roundtrip(self, tmp_path: Path, spacy_only_engine):
        original = "客戶身分證A123456789，手機0912345678"
        f = tmp_path / "data.txt"
        f.write_text(original, encoding="utf-8")

        content = read_file(f)
        anonymized_text, mapping = spacy_only_engine.anonymize(content.text)
        assert "A123456789" not in anonymized_text

        from pii_guard.pipeline.engine import PiiGuardEngine

        restored = PiiGuardEngine.deanonymize(anonymized_text, mapping)
        assert restored == original

    def test_csv_roundtrip(self, tmp_path: Path, spacy_only_engine):
        original = "name,phone\n張三,0912345678\n李四,0923456789"
        f = tmp_path / "data.csv"
        # write_bytes, not write_text: write_text translates newlines on
        # Windows, so the on-disk file would not be what this test compares
        # the round-trip against.
        f.write_bytes(original.encode("utf-8"))

        content = read_file(f)
        anonymized_text, mapping = spacy_only_engine.anonymize(content.text)
        assert "0912345678" not in anonymized_text

        from pii_guard.pipeline.engine import PiiGuardEngine

        restored = PiiGuardEngine.deanonymize(anonymized_text, mapping)
        assert restored == original

    def test_xlsx_roundtrip(self, tmp_path: Path, spacy_only_engine):
        import openpyxl

        # Create
        wb = openpyxl.Workbook()
        ws = wb.active
        ws["A1"] = "身分證"
        ws["B1"] = "A123456789"
        src = tmp_path / "data.xlsx"
        wb.save(str(src))
        wb.close()

        # Anonymize
        content = read_file(src)
        anonymized_text, mapping = spacy_only_engine.anonymize(content.text)
        anon_out = tmp_path / "anon.xlsx"
        write_file(content, anonymized_text, mapping, anon_out)

        # Verify anonymized xlsx
        wb2 = openpyxl.load_workbook(str(anon_out), read_only=True)
        ws2 = wb2.active
        assert "A123456789" not in str(ws2["B1"].value)
        wb2.close()
